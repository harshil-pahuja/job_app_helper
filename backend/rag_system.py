"""
RAG (Retrieval-Augmented Generation) system for resume and job description processing.
Handles embeddings, vector storage, and document retrieval.
"""
import json
import os
import logging
import re
from dotenv import load_dotenv
from langchain_community.vectorstores import Chroma  # type: ignore
from langchain_core.documents import Document  # type: ignore
from langchain_core.messages import HumanMessage, SystemMessage  # type: ignore
from langchain_core.tools import tool  # type: ignore
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings  # type: ignore
from langchain_community.retrievers import BM25Retriever  # type: ignore
from sentence_transformers import CrossEncoder  # type: ignore

load_dotenv()
os.environ.setdefault("TIKTOKEN_CACHE_DIR", os.path.join(os.getcwd(), ".chroma_test"))

logger = logging.getLogger(__name__)
DEBUG_PRIVACY_LOGS = os.getenv("DEBUG_PRIVACY_LOGS", "").lower() == "true"


def quiet_third_party_http_logs():
    """Keep local debug output readable by hiding dependency HTTP chatter."""
    for logger_name in (
        "httpx",
        "httpcore",
        "huggingface_hub",
        "sentence_transformers",
        "urllib3",
    ):
        logging.getLogger(logger_name).setLevel(logging.WARNING)


class RAGSystem:
    def __init__(
        self,
        collection_name: str = "job_app_helper_collection",
        embedding_backend: str = "auto",
        persist_directory: str | None = None,
    ):
        self.collection_name = collection_name
        self.persist_directory = persist_directory
        self.embeddings = self._build_embeddings(embedding_backend)
        self.vectorstore = None
        self.documents: list[Document] = []
        self.section_aliases = self._get_section_aliases()
        self.section_headers = [
            alias
            for aliases in self.section_aliases.values()
            for alias in aliases
        ]
        self.alias_to_section = self._build_alias_to_section_map(self.section_aliases)
        self.section_pattern = self._build_section_pattern(self.section_headers)
        self.header_classification_cache: dict[str, str | None] = {}
        self.horizontal_rule_header_candidates: set[str] = set()
        self.cross_encoder_model = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")

    def _build_embeddings(self, embedding_backend: str):
        """Build embedding model. Auto mode prefers OpenAI if API key exists."""
        backend = embedding_backend.lower().strip()
        if backend == "auto":
            backend = "openai" if os.getenv("OPENAI_API_KEY") else "huggingface"

        if backend == "openai":
            return OpenAIEmbeddings()

        if backend == "huggingface":
            # Local embedding model for offline/low-cost testing.
            return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

        raise ValueError("embedding_backend must be one of: auto, openai, huggingface")

    def extract_and_read_resume_text(self, resume_file_path: str) -> str:
        """Extract text from a resume file (PDF, DOCX, or TXT)."""
        from pathlib import Path
        import pdfplumber #type: ignore

        path = Path(resume_file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {resume_file_path}")

        if path.suffix.lower() == ".pdf":
            with pdfplumber.open(str(path)) as pdf:
                page_texts = []
                for page in pdf.pages:
                    text = page.extract_text(
                        layout=True,
                        x_tolerance=1,
                        y_tolerance=3,
                    )
                    if not text:
                        text = page.extract_text() or ""
                    page_texts.append(text)
                return "\n".join(page_texts).strip()

        if path.suffix.lower() == ".docx" or path.suffix.lower() == ".doc":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n".join((para.text for para in doc.paragraphs)).strip()

        if path.suffix.lower() == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore").strip()

        raise ValueError("Unsupported resume file format. Supported formats: PDF, DOCX, TXT.")
    
    def handle_multi_word_headers(self, text: str) -> str:
        """Repair known multi-word section headers that are collapsed into body text."""
        import re

        repaired_text = text
        for header in sorted(self.section_headers, key=len, reverse=True):
            escaped_header = re.escape(header)
            collapsed_header_pattern = (
                rf"(?im)^(\s*{escaped_header}:)\s+"
                rf"|^(\s*{escaped_header})\s+(?=([*\-\u2022\u25cf]|â\S*)\s*\S)"
                rf"|^(\s*{escaped_header})\s+(?=[A-Z0-9])"
            )
            repaired_text = re.sub(
                collapsed_header_pattern,
                lambda match: f"{(match.group(1) or match.group(2) or match.group(4)).strip()}\n",
                repaired_text,
            )
        return repaired_text

    def cleanup_horizontal_rule_headers(self, text: str) -> str:
        """Remove PDF underline artifacts while preserving the header text before them."""
        import re

        cleaned_lines = []
        self.horizontal_rule_header_candidates = set()
        horizontal_rule_chars = r"_\-=\u2012\u2013\u2014\u2015\u2500\u2501"
        rule_only_pattern = re.compile(rf"^[\s{horizontal_rule_chars}]{{5,}}$")
        trailing_rule_pattern = re.compile(
            rf"^(?P<header>.+?)\s+[{horizontal_rule_chars}]{{5,}}\s*$"
        )

        lines = text.splitlines()
        for index, line in enumerate(lines):
            stripped_line = line.strip()
            following_lines = lines[index + 1:]
            has_following_content = any(
                candidate.strip()
                and not rule_only_pattern.fullmatch(candidate.strip())
                for candidate in following_lines
            )

            if rule_only_pattern.fullmatch(stripped_line):
                if has_following_content:
                    if cleaned_lines:
                        previous_line = cleaned_lines[-1].strip()
                        if previous_line:
                            self.horizontal_rule_header_candidates.add(
                                previous_line.rstrip(":").strip().lower()
                            )
                    continue
                cleaned_lines.append(line)
                continue

            trailing_rule_match = trailing_rule_pattern.match(stripped_line)
            if trailing_rule_match and has_following_content:
                header = trailing_rule_match.group("header").strip()
                self.horizontal_rule_header_candidates.add(
                    header.rstrip(":").strip().lower()
                )
                cleaned_lines.append(header)
                continue

            cleaned_lines.append(line)

        return "\n".join(cleaned_lines)
    
    def normalize_resume_text(self, text: str) -> str:
        """Normalize resume text for consistent section detection."""
        import re
        # Replace multiple spaces/tabs with a single space
        normalized_text = re.sub(r'[ \t\f\v]+', ' ', text).strip()
        # Remove PDF horizontal-rule artifacts while keeping header text.
        normalized_text = self.cleanup_horizontal_rule_headers(normalized_text)
        # Collapse multiple newlines into two newlines
        normalized_text = re.sub(r'\n{3,}', '\n\n', normalized_text)
        # Repair collapsed multi-word section headers
        normalized_text = self.handle_multi_word_headers(normalized_text)
        return normalized_text

    def classify_header_candidate(self, line: str, section_headers: list[str]) -> str:
        """Classify a line as not_header, subsection, or top_level."""
        import re

        line = line.strip()
        if not line:
            return "not_header"

        header_text = line.rstrip(":").strip()
        words = header_text.split()
        if len(line) > 80 or len(words) > 6:
            return "not_header"
        if re.search(r"^\s*[-*â€¢]\s+", line):
            return "not_header"
        if re.search(r"[.!?;]$", line):
            return "not_header"
        if "@" in line or "http" in line.lower() or "www." in line.lower():
            return "not_header"
        if len(re.findall(r"\d", line)) > 2:
            return "not_header"

        action_verbs = {
            "built", "created", "developed", "implemented", "improved",
            "managed", "led", "designed", "collaborated", "analyzed",
            "graduated", "available", "worked", "used",
        }
        first_word = words[0].lower().strip(":") if words else ""
        if first_word in action_verbs:
            return "not_header"

        if not re.fullmatch(r"[A-Za-z][A-Za-z0-9 &/+\-,:]*", line):
            return "not_header"

        normalized_line = header_text.lower()
        is_known_header = any(normalized_line == header.lower() for header in section_headers)
        is_top_level_format = header_text.isupper() or (not line.endswith(":") and header_text.istitle())

        if line.endswith(":"):
            return "top_level" if header_text.isupper() else "subsection"

        if is_known_header and is_top_level_format:
            return "top_level"

        if header_text.isupper() or header_text.istitle():
            return "top_level"

        if len(words) <= 3 and all(len(word) > 2 for word in words):
            return "subsection"

        return "not_header"

    def _classify_compound_header_with_llm(self, header: str) -> str | None:
        """Classify compound resume headers containing connectors such as and, /, or &."""
        import re

        if not self._has_multi_word_header_operator(header):
            return None

        normalized_header = header.strip().lower()
        if normalized_header in self.header_classification_cache:
            return self.header_classification_cache[normalized_header]

        allowed_sections = list(self.section_aliases.keys()) + ["General"]
        if not os.getenv("OPENAI_API_KEY"):
            self.header_classification_cache[normalized_header] = None
            return None

        try:
            model = ChatOpenAI(
                model=os.getenv("RAG_HEADER_CLASSIFIER_MODEL", "gpt-5.6-terra"),
                temperature=0,
                timeout=20,
                max_retries=1,
            )
            response = model.invoke(
                [
                    SystemMessage(
                        content=(
                            "You classify compound resume section headers that use connectors "
                            "such as 'and', '&', '/', '|', or '+'. "
                            "Return only valid JSON. Do not include explanation or markdown."
                        )
                    ),
                    HumanMessage(
                        content=(
                            "Map this compound resume header to exactly one canonical section.\n"
                            "Only classify the line if it is a resume section header formed with "
                            "connectors/operators like 'and', '&', '/', '|', or '+'.\n\n"
                            "Do NOT break apart text before and after connectors into multiple chunks. Only return one chunk with the right section.\n"
                            "Example: If a resume header is 'Leadership and Management', map it to 'Leadership' "
                            "but do not break the words Leadership and Management into two chunks.\n\n"
                            f"Allowed sections: {', '.join(allowed_sections)}\n"
                            f"Header: {header!r}\n\n"
                            'Return JSON like {"section": "Leadership"}. '
                            'Use "General" if the line is not a resume section header.' 
                        )
                    ),
                ]
            )
            raw_content = re.sub(
                r"^```(?:json)?\s*|\s*```$",
                "",
                str(response.content).strip(),
                flags=re.IGNORECASE,
            )
            json_match = re.search(r"\{.*\}", raw_content, re.DOTALL)
            parsed = json.loads(json_match.group() if json_match else raw_content)
            section = parsed.get("section")
            if section in self.section_aliases:
                self.header_classification_cache[normalized_header] = section
                return section
            self.header_classification_cache[normalized_header] = None
            return None
        except Exception as exc:
            logger.debug("LLM header classification failed for %r: %s", header, exc)
            self.header_classification_cache[normalized_header] = None
            return None

    def _has_multi_word_header_operator(self, header: str) -> bool:
        """Return whether a line has connectors that warrant LLM header classification."""
        import re

        return re.search(r"^(?:and|or)\b|\s+(?:and|or)\s+|[&/|+]", header, re.IGNORECASE) is not None

    def _starts_with_header_connector(self, line: str) -> bool:
        """Return whether a line starts as a continuation of a previous header."""
        import re

        return re.search(r"^(?:(?:and|or)\b|[&/|+])", line.strip(), re.IGNORECASE) is not None

    def _get_default_section_aliases(self) -> dict[str, list[str]]:
        """Return fallback canonical resume section labels and known aliases."""
        return {
            "Education": [
                "Education", "Academic Background", "Academic Qualifications",
                "Academic Training", "Academic Journey",
            ],
            "Skills": [
                "Skills", "Technical Skills", "Soft Skills", "Core Competencies",
                "Competencies", "Technical Expertise", "Areas of Expertise",
                "Key Skills", "Technical Proficiencies", "Tools & Technologies",
                "Tools and Technologies",
            ],
            "Experience": [
                "Experience", "Experiences", "Work Experience", "Work Experiences",
                "Professional Experience", "Professional Background", "Employment",
                "Employment History", "Work History", "Career History",
                "Relevant Experience",
            ],
            "Projects": [
                "Project", "Projects", "Personal Projects", "Academic Projects",
                "Key Projects", "Selected Projects", "Selected Work",
            ],
            "Leadership": [
                "Leadership", "Leadership Experience", "Activities", "Involvement", "Extracurriculars"
            ],
            "Certifications": [
                "Certifications", "Licenses", "Certificates",
            ],
            "Awards": [
                "Awards", "Honors", "Achievements",
                "Accomplishments",
            ],
            "Summary": [
                "Summary", "Professional Summary", "Executive Summary",
                "Objective", "Career Objective", "Profile",
            ],
            "Publications": ["Publications"],
            "Volunteer": ["Volunteer", "Community Involvement"],
            "References": ["References"],
            "Contact": ["Contact"],
        }

    def _get_section_aliases(self) -> dict[str, list[str]]:
        """Return resume section aliases from the LLM, falling back to defaults."""
        default_aliases = self._get_default_section_aliases()
        canonical_sections = list(default_aliases.keys())

        try:
            model = ChatOpenAI(model="gpt-5.6-terra", temperature=0, timeout=30, max_retries=1)
            response = model.invoke(
                [
                    SystemMessage(
                        content=(
                            "You generate robust resume section-header aliases for a RAG chunking system. "
                            "Return only valid JSON. Do not include explanations. "
                            "The JSON object must use exactly the provided canonical section labels as keys. "
                            "Each value must be a list of realistic resume header strings that should map to that section."
                        )
                    ),
                    HumanMessage(
                        content=(
                            "Canonical section labels:\n"
                            f"{json.dumps(canonical_sections)}\n\n"
                            "Generate aliases that cover common resume formats, including singular/plural forms, "
                            "academic/professional wording, compound headers, and common synonyms. "
                            "Return JSON only."
                            "Examples: {\"Education\": [\"Education\", \"Academic Background\", \"Academic Qualifications\"]}, " \
                            "\"Skills\": [\"Skills\", \"Technical Skills\", \"Soft Skills\"]}"
                        )
                    ),
                ]
            )
            json_match = re.search(r"\{.*\}", response.content, re.DOTALL)
            if not json_match:
                return default_aliases

            parsed = json.loads(json_match.group())
            if not isinstance(parsed, dict):
                return default_aliases

            aliases_by_section: dict[str, list[str]] = {}
            for section in canonical_sections:
                llm_aliases = parsed.get(section, [])
                if not isinstance(llm_aliases, list):
                    llm_aliases = []

                cleaned_aliases = []
                for alias in llm_aliases:
                    if isinstance(alias, str):
                        cleaned_alias = re.sub(r"\s+", " ", alias).strip()
                        if cleaned_alias:
                            cleaned_aliases.append(cleaned_alias)

                aliases = [section] + cleaned_aliases + default_aliases[section]
                aliases_by_section[section] = list(dict.fromkeys(aliases))

            return aliases_by_section
        except Exception as exc:
            logger.debug("LLM section alias generation failed; using defaults. Error: %s", exc)
            return default_aliases

    def _build_alias_to_section_map(self, section_aliases: dict[str, list[str]]) -> dict[str, str]:
        """Map each lowercase alias to its canonical section label."""
        return {
            alias.lower(): section
            for section, aliases in section_aliases.items()
            for alias in aliases
        }

    def cross_encoder_rerank(self, query: str, candidate_chunks: list[Document], top_k: int = 3) -> list[Document]:
        """Rerank candidate chunks using a cross-encoder model for better relevance."""
        pairs = [(query, chunk.page_content) for chunk in candidate_chunks]
        relevance_scores = self.cross_encoder_model.predict(pairs)
        scored_chunks = sorted(list(zip(candidate_chunks, relevance_scores)), reverse=True, key=lambda x: x[1])
        return [chunk for chunk, _ in scored_chunks[:top_k]]

    def _build_section_pattern(self, section_headers: list[str]) -> str:
        """Build the escaped regex pattern for known section headers."""
        import re
        return "|".join([re.escape(header) for header in section_headers])

    def _match_compound_header_alias(self, line: str) -> str | None:
        """Map compound top-level headers to a canonical section using known aliases."""
        import re

        if not self._has_multi_word_header_operator(line):
            return None

        normalized_line = re.sub(r"[^a-z0-9]+", " ", line.lower()).strip()
        if not normalized_line:
            return None

        for section, aliases in self.section_aliases.items():
            for alias in sorted(aliases, key=len, reverse=True):
                normalized_alias = re.sub(r"[^a-z0-9]+", " ", alias.lower()).strip()
                if not normalized_alias:
                    continue
                if re.search(rf"\b{re.escape(normalized_alias)}\b", normalized_line):
                    return section

        return None

    def classify_compound_header(self, line: str) -> str | None:
        """Classify a compound resume section header using aliases first, then the LLM."""
        if not self._has_multi_word_header_operator(line):
            return None

        alias_section = self._match_compound_header_alias(line)
        if alias_section:
            return alias_section

        return self._classify_compound_header_with_llm(line)

    def _is_section_boundary_line(
        self,
        line: str,
        section_headers: list[str],
        alias_to_section: dict[str, str],
    ) -> bool:
        """Return whether a standalone line should start a new resume section chunk."""
        import re

        if self._starts_with_header_connector(line):
            return False

        normalized_line = line.rstrip(":").strip().lower()
        candidate_type = self.classify_header_candidate(line, section_headers)
        if (
            normalized_line in getattr(self, "horizontal_rule_header_candidates", set())
            and candidate_type == "top_level"
        ):
            return True

        if normalized_line in alias_to_section and candidate_type == "top_level":
            return True

        exact_header = any(
            re.search(f'(?i)^\\s*{re.escape(header)}\\s*:?\\s*$', line)
            for header in section_headers
        )
        if exact_header and candidate_type == "top_level":
            return True

        if candidate_type != "top_level":
            return False

        if self.classify_compound_header(line):
            return True

        return False

    def _should_merge_header_continuation(
        self,
        previous_line: str,
        current_line: str,
        section_headers: list[str],
    ) -> bool:
        """Return whether current_line should be merged into the previous header line."""
        import re

        previous_line = previous_line.strip()
        current_line = current_line.strip()
        if not previous_line or not current_line:
            return False

        if not self._starts_with_header_connector(current_line):
            return False

        if self.classify_header_candidate(previous_line, section_headers) == "not_header":
            return False

        combined_header = f"{previous_line} {current_line}"
        if self.classify_compound_header(combined_header):
            return True

        return previous_line.rstrip(":").strip().lower() in self.alias_to_section

    def _label_chunk_section(
        self,
        first_line: str,
        section_headers: list[str],
        section_aliases: dict[str, list[str]],
        alias_to_section: dict[str, str],
    ) -> str:
        """Return the canonical section label for a chunk's first line."""
        import re

        matched_header = next(
            (header for header in section_headers if re.search(f'(?i)^\\s*{re.escape(header)}\\s*:?\\s*$', first_line)),
            None,
        )
        if matched_header:
            return alias_to_section.get(matched_header.lower(), matched_header)

        if self.classify_header_candidate(first_line, section_headers) != "not_header":
            compound_section = self.classify_compound_header(first_line)
            if compound_section:
                return compound_section

        return "General"

    def _fallback_chunk_text(
        self,
        normalized_text: str,
        chunk_size: int = 1000,
        overlap: int = 200,
    ) -> list[tuple[str, str]]:
        """Fallback chunking when no reliable section headers are detected."""
        chunks = []
        step = max(chunk_size - overlap, 1)
        for i in range(0, len(normalized_text), step):
            chunk = normalized_text[i:i + chunk_size]
            if chunk.strip():
                chunks.append((chunk, "General"))
        return chunks
    
    def _split_text_into_chunks_with_sections(
        self,
        text: str,
        section_headers: list[str] | None = None,
        chunk_size: int = 1000,
        overlap: int = 200,
    ) -> list[tuple[str, str]]:
        """Split text into chunks while tracking which resume section each chunk belongs to."""
        section_aliases = self.section_aliases
        alias_to_section = self.alias_to_section
        section_headers = section_headers or self.section_headers
        section_pattern = self.section_pattern if section_headers is self.section_headers else self._build_section_pattern(section_headers)
        logger.debug("Section pattern: %s", section_pattern)

        normalized_text = self.normalize_resume_text(text)
        logger.debug("Normalized text preview: %s...", normalized_text[:100])

        if not normalized_text:
            return []

        section_texts = []
        current_lines = []
        found_section_boundary = False

        for line in normalized_text.splitlines():
            stripped_line = line.strip()
            if (
                current_lines
                and self._should_merge_header_continuation(
                    current_lines[-1],
                    stripped_line,
                    section_headers,
                )
            ):
                current_lines[-1] = current_lines[-1].rstrip() + " " + stripped_line
                continue

            is_section_boundary = bool(stripped_line) and self._is_section_boundary_line(
                stripped_line,
                section_headers,
                alias_to_section,
            )
            if is_section_boundary:
                found_section_boundary = True

            if current_lines and is_section_boundary:
                section_texts.append("\n".join(current_lines).strip())
                current_lines = [line]
            else:
                current_lines.append(line)

        if current_lines:
            section_texts.append("\n".join(current_lines).strip())

        if not found_section_boundary:
            logger.debug("No section headers found in text. Defaulting to general chunking.")
            return self._fallback_chunk_text(normalized_text, chunk_size=chunk_size, overlap=overlap)
                
        chunks = []
        for section_text in section_texts:
            if not section_text:
                continue
            first_line = section_text.splitlines()[0].strip()
            section_header = self._label_chunk_section(
                first_line,
                section_headers,
                section_aliases,
                alias_to_section,
            )
            logger.debug("Chunk preview: %s... | Section: %s", section_text[:30], section_header)
            chunks.append((section_text, section_header))

        return chunks

    def load_and_process_document(
        self,
        file_path: str,
    ) -> list[tuple[str, str]]:
        """Load a document, extract text, normalize, and split into chunks with section labels."""
        text = self.extract_and_read_resume_text(file_path)
        if not text:
            raise ValueError("No text extracted from the document.")
        chunks = self._split_text_into_chunks_with_sections(text)
        section_labels = [section for _, section in chunks]
        logger.debug("Section labels for chunks: %s", section_labels)
        return chunks


    def create_vectorstore(self, chunks):
        """Create a vector store from the document chunks."""
        valid_chunks = []
        for index, chunk in enumerate(chunks):
            if isinstance(chunk, Document):
                if chunk.page_content and chunk.page_content.strip():
                    chunk.metadata.setdefault("chunk_index", index)
                    valid_chunks.append(chunk)
                continue

            content, section = chunk
            if content and content.strip():
                valid_chunks.append(
                    Document(
                        page_content=content,
                        metadata={"chunk_index": index, "section": section},
                    )
                )

        logger.debug("Valid chunks: %d", len(valid_chunks))

        if not valid_chunks:
            raise ValueError("No valid chunks to create vectorstore (all chunks are empty)")

        self.documents = valid_chunks
        self.vectorstore = Chroma.from_documents(
            documents=valid_chunks,
            embedding=self.embeddings,
            collection_name=self.collection_name,
            persist_directory=self.persist_directory,
        )

    def retrieve_relevant_chunks(self, query: str, top_k: int = 3):
        """Retrieve relevant document chunks based on a query, sorted by document order."""
        if not self.vectorstore:
            raise ValueError("Vector store not created. Please load and process a document first.")

        #Access all chunks for RRF scoring, then filter to top_k after sorting by document order
        candidate_k = len(self.documents)
        bm25_retriever = BM25Retriever.from_documents(self.documents, k=candidate_k)
        bm25_results = bm25_retriever.invoke(query)
        vector_results = self.vectorstore.similarity_search(query=query, k=candidate_k)

        rrf_results_for_bm25 = []  # Stores key value pairs of (Document, RRF score) for BM25 results
        for rank, key in enumerate(bm25_results, start=1):
            rrf_score = 1 / (60 + rank)  #60 is a constant k used in production
            rrf_results_for_bm25.append((key, rrf_score))

        rrf_results_for_vector = []
        for rank, key in enumerate(vector_results, start=1):
            rrf_score = 1 / (60 + rank) #60 is a constant k used in production
            rrf_results_for_vector.append((key, rrf_score))

        combined_results: dict[int, tuple[Document, float]] = {}
        all_rrf_results = rrf_results_for_bm25 + rrf_results_for_vector
        for doc, score in all_rrf_results:
            chunk_index = doc.metadata.get("chunk_index")
            if chunk_index is not None:
                if chunk_index in combined_results:
                    existing_doc, existing_score = combined_results[chunk_index]
                    combined_results[chunk_index] = (existing_doc, existing_score + score)
                else:
                    combined_results[chunk_index] = (doc, score)

        ranked_chunks = [doc for doc, _ in sorted(combined_results.values(), key=lambda x: x[1], reverse=True)]
        return self.cross_encoder_rerank(query, ranked_chunks, top_k=top_k)

    def ingest_file(self, file_path: str):
        """Convenience method for tests: load, split, and index a file in one call."""
        chunks = self.load_and_process_document(file_path)
        self.create_vectorstore(chunks)

def create_retrieve_resume_tool(rag_instance: RAGSystem):
    """Factory function to create a retrieval tool bound to a specific RAG instance.
    
    Args:
        rag_instance: The RAGSystem instance to retrieve from
        
    Returns:
        A LangChain tool function the agent can call
    """
    @tool
    def retrieve_resume_context(query: str) -> str:
        """Retrieve relevant resume sections for the given query."""
        chunks = rag_instance.retrieve_relevant_chunks(query, top_k=3)
        print(f"Top chunk retrieved for query '{query}': {chunks[0].page_content[:100]}..." if chunks else "No chunks retrieved.")
        if not chunks:
            return "No relevant information found in the resume."
        return "\n\n".join([chunk.page_content for chunk in chunks])
    
    return retrieve_resume_context

def main():
    """Test selected resumes with comparison-oriented retrieval tasks."""
    from pathlib import Path

    if not logger.handlers:
        logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    quiet_third_party_http_logs()

    resume_dir = Path("tests/resumes")
    retrieval_tests = [
        {
            "task": "education_evidence",
            "query": "degree major field of study university coursework education bachelor",
            "expected_sections": {"Education"},
        },
        {
            "task": "skills_evidence",
            "query": "programming languages frameworks tools technologies technical skills",
            "expected_sections": {"Skills", "Projects", "Experience"},
        },
        {
            "task": "project_evidence",
            "query": "projects applications websites software systems implementation development",
            "expected_sections": {"Projects"},
        },
        {
            "task": "leadership_evidence",
            "query": "leadership extracurriculars organizations activities volunteer clubs involvement",
            "expected_sections": {"Leadership"},
        },
        {
            "task": "experience_evidence",
            "query": "work experience employers companies roles responsibilities internship professional experience",
            "expected_sections": {"Experience", "Leadership"},
        },
        {
            "task": "certification_or_awards_evidence",
            "query": "certifications certificates awards honors scholarships coursework programs",
            "expected_sections": {"Certifications", "Education", "Leadership"},
        },
    ]
    resume_tests = [
        {"name": "Isabella", "resume_match": "Isabella"},
    ]

    def find_resume(match_text: str) -> Path | None:
        supported_suffixes = {".pdf", ".txt", ".doc", ".docx"}
        normalized_match = match_text.lower()
        for path in sorted(resume_dir.iterdir()):
            if (
                path.is_file()
                and path.suffix.lower() in supported_suffixes
                and normalized_match in path.name.lower()
            ):
                return path
        return None

    logger.info("=" * 70)
    logger.info("Testing %d resumes with comparison-oriented RAG retrieval", len(resume_tests))

    for resume_index, resume_test in enumerate(resume_tests, start=1):
        resume_path = find_resume(resume_test["resume_match"])
        logger.info("\n" + "=" * 70)
        logger.info("Resume %d/%d: %s", resume_index, len(resume_tests), resume_test["name"])
        logger.info("Resume match: %s", resume_test["resume_match"])

        if not resume_path:
            logger.error("Resume not found in %s", resume_dir)
            continue

        logger.info("Resume: %s", resume_path.name)
        safe_name = re.sub(r"[^a-zA-Z0-9._-]+", "_", resume_test["name"].lower()).strip("._-")
        rag = RAGSystem(collection_name=f"debug_{safe_name}_rag")

        try:
            chunks = rag.load_and_process_document(str(resume_path))
        except Exception as exc:
            logger.error("Structure-aware chunking failed. Error: %s", exc)
            continue

        logger.info("\nStructure-aware chunks")
        for chunk_index, (chunk_text, section) in enumerate(chunks):
            first_line = chunk_text.splitlines()[0].strip() if chunk_text.splitlines() else ""
            logger.info(
                "- chunk_index=%s section=%s first_line=%r chars=%d",
                chunk_index,
                section,
                first_line,
                len(chunk_text),
            )

        try:
            rag.create_vectorstore(chunks)
        except Exception as exc:
            logger.error("Vectorstore creation failed. Error: %s", exc)
            continue

        logger.info("\nRetrieval evidence tests")
        for test_index, test in enumerate(retrieval_tests, start=1):
            logger.info("\n" + "-" * 70)
            logger.info("Task %d/%d: %s", test_index, len(retrieval_tests), test["task"])
            logger.info("Query: %s", test["query"])
            logger.info("Expected sections: %s", sorted(test["expected_sections"]))

            try:
                relevant_chunks = rag.retrieve_relevant_chunks(test["query"], top_k=3)
            except Exception as exc:
                logger.error("Retrieval failed. Error: %s", exc)
                continue

            returned_sections = [
                chunk.metadata.get("section", "Unknown")
                for chunk in relevant_chunks
            ]

            status = "PASS" if any(section in test["expected_sections"] for section in returned_sections) else "REVIEW"
            logger.info("Returned sections: %s", returned_sections)
            logger.info("Status: %s", status)

            for result_index, chunk in enumerate(relevant_chunks, start=1):
                chunk_lines = chunk.page_content.splitlines()
                first_line = chunk_lines[0].strip() if chunk_lines else ""
                preview = " ".join(chunk.page_content.split())[:350]
                logger.info(
                    "- result=%s chunk_index=%s section=%s first_line=%r chars=%d",
                    result_index,
                    chunk.metadata.get("chunk_index"),
                    chunk.metadata.get("section"),
                    first_line,
                    len(chunk.page_content),
                )
                logger.info("  evidence preview: %s", preview)

    logger.info("=" * 70)


if __name__ == "__main__":
    main()
